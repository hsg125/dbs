from docx import Document
import string
import random
import win32com
import os
from win32com.client import Dispatch

KEY_LEN = 20


# 产生随机文件名
def base_str():
    return string.ascii_letters + string.digits


def key_gen():
    keylist = [random.choice(base_str()) for i in range(KEY_LEN)]
    return "".join(keylist)


# 删除临时文件
def del_file(file_name):
    try:
        os.remove('.\\temp\\' + file_name + '.docx')
    except IOError as e:
        err = str(e) + '\n系统错误，无法删除文件-' + file_name + '，可能被占用'
        return err


# 打开文档
def open_word(filename):
    try:
        doc = Document('.\\templates\\' + filename + '.docx')
        return doc
    except Exception as e:
        err = str(e) + '\n无法打开文件，请重启后尝试。'
        return err


# 牙位转换
def trans(pos):
    # pos = ['322','','','']
    p = []
    for i in range(4):
        if len(pos[i]) > 0:
            for x in pos[i]:
                p.append(str(i+1) + x)
    if not len(p):
        s = '未指定牙位'
    else:
        s = '.'.join(p)
    return s


# 修改文档内容，并存为随机文件
def get_word(filename, s):
    doc = open_word(filename)
    if str(type(doc)) == "<class 'str'>":
        return doc
    tab = doc.tables
    tab0 = tab[0]
    cells = tab0._cells
    text = ''
    # 判断是否为加工单
    if filename == 'jgd':
        cells[8].paragraphs[0].runs[0].text = s[1]  # 医生
        cells[16].paragraphs[0].runs[0].text = s[2]  # 患者
        cells[19].paragraphs[0].runs[0].text = '13866676217'  # 电话
        pos = eval(s[-2])
        cells[34].paragraphs[0].runs[0].text = pos[0]  # lu
        cells[36].paragraphs[0].runs[0].text = ''  # -
        cells[37].paragraphs[0].runs[0].text = pos[1]  # ru
        cells[42].paragraphs[0].runs[0].text = pos[2]  # ld
        cells[44].paragraphs[0].runs[0].text = ''  # -
        cells[45].paragraphs[0].runs[0].text = pos[3]  # rd
        if s[4] == '':
            text = s[3] + s[-1]
        else:
            text = s[3] + '\n颜色：' + s[4] + '\n' + s[-1]
        cells[57].paragraphs[0].text = text  # text
    # 判断是修复知情同意书
    elif filename == 'gdxf' or filename == 'hdxf':
        # type = ['拔除术','根管治疗术', '牙冠延长术', '牙周翻瓣术','固定修复术','活动修复术']
        s[1] = (lambda x, y: x if s[1] == 3 else y)('固定修复术', '活动修复术')
        s[3] = (lambda x, y: x if s[3] == 0 else y)('女', '男')

        # cells_str = [cell.text for cell in cells]
        # print(cells_str)
        header = doc.sections[0].header
        header.paragraphs[0].text += str(s[-1]).rjust(6, '0')

        cells[0].paragraphs[0].runs[0].text += s[2]  # 患者姓名
        cells[1].paragraphs[0].runs[1].text += s[3]  # 性别
        cells[2].paragraphs[0].runs[1].text += str(s[4])  # 年龄
        cells[3].paragraphs[0].runs[1].text += '口腔科'  # 科别
        cells[4].paragraphs[0].runs[1].text += str(s[5])  # 病历号
        pos = eval(s[8])
        ps = trans(pos)
        cells[5].paragraphs[1].runs[0].text = cells[5].paragraphs[1].runs[0].text.replace('告知我患有：',
                                                                                          '  ' + '告知我患有：' + s[7] + '  ')
        cells[5].paragraphs[2].runs[3].text = cells[5].paragraphs[2].runs[3].text.replace('牙位',
                                                                                          '  ' + ps + '  ')
        # cells[5].paragraphs[2].runs[2].text = cells[5].paragraphs[2].runs[2].text.replace(source[2],
        #                                                                                   '  ' + source[2] + dst[2] + '  ')
        cells[25].paragraphs[6].runs[4].text += '胡韶光'
        cells[25].paragraphs[8].runs[5].text += str(s[0])

    else:
        type1 = ['拔除术', '根管治疗术', '牙冠延长术', '牙周翻瓣术']
        if s[1] == 0:
            s[1] = type1[0]
        elif s[1] == 1:
            s[1] = type1[1]
        elif s[1] == 2 and s[7] == '慢性牙周炎':
            s[1] = type1[3]
        else:
            s[1] = type1[2]

        # s[1] = (lambda x, y: x if s[1] == 3 else y)('固定修复术', '活动修复术')
        s[3] = (lambda x, y: x if s[3] == 0 else y)('女', '男')

        # cells_str = [cell.text for cell in cells]
        # print(cells_str)
        header = doc.sections[0].header
        header.paragraphs[0].text += str(s[-1]).rjust(6, '0')

        cells[0].paragraphs[0].runs[0].text += s[2]  # 患者姓名
        cells[1].paragraphs[0].runs[1].text += s[3]  # 性别
        cells[2].paragraphs[0].runs[1].text += str(s[4])  # 年龄
        cells[3].paragraphs[0].runs[1].text += '口腔科'  # 科别
        cells[4].paragraphs[0].runs[1].text += str(s[5])  # 病历号
        pos = eval(s[8])
        ps = trans(pos)

        cells[5].paragraphs[1].runs[3].text = cells[5].paragraphs[1].runs[3].text.replace('疾病名称',
                                                                                          '  ' + ps + s[7] + '  ')
        cells[5].paragraphs[2].runs[3].text = cells[5].paragraphs[2].runs[3].text.replace('麻醉方法',
                                                                                          '  ' + s[6] + '  ')
        cells[5].paragraphs[2].runs[7].text = cells[5].paragraphs[2].runs[7].text.replace('手术名称',
                                                                                          '  ' + ps + s[1] + '  ')
        if s[1] == '根管治疗术':
            cells[25].paragraphs[5].runs[4].text += '胡韶光'
            cells[25].paragraphs[7].runs[2].text += str(s[0])
        else:
            cells[25].paragraphs[6].runs[4].text += '胡韶光'
            cells[25].paragraphs[8].runs[2].text += str(s[0])

    random_filename = key_gen()
    doc.save('.\\temp\\' + random_filename + '.docx')
    return random_filename


# 通过win32com的方式进行打印
def print_word(word_file):
    exec_tool = 'Word.Application'
    # 如果使用wps
    # exec_tool = 'wps.application'
    word = win32com.client.Dispatch(exec_tool)
    # 在后台运行程序
    word.Visible = 0  # 后台运行，不显示
    # 运行过程不警告
    word.DisplayAlerts = 0  # 不警告
    # 打开word文档
    path = os.path.realpath('./') + '\\'
    doc = word.Documents.Open(path + '\\temp\\' + word_file + '.docx')

    # 进行打印

    doc.PrintOut()
    # --------------------------------------------------------
    # -- 最后操作保存
    # --------------------------------------------------------
    # 关闭文件
    doc.Close()
    # 退出word
    word.Quit()

# --------------------------------------
